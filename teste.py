import os
import re
import hmac
import json
import time
import hashlib
from datetime import datetime
from typing import Optional, Tuple
from threading import Thread

import requests
from fastapi import FastAPI, Request, HTTPException
from fastapi.responses import JSONResponse

from PIL import Image, ImageEnhance, ImageFilter, ImageDraw, ImageFont
import pytesseract

import fitz  # PyMuPDF

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

try:
    from google.oauth2.service_account import Credentials
    from google.oauth2.credentials import Credentials as OAuthCredentials
    from google.auth.transport.requests import Request as GoogleAuthRequest
    from googleapiclient.discovery import build
    GSHEETS_AVAILABLE = True
except ImportError:
    GSHEETS_AVAILABLE = False

# =========================
# Config
# =========================

app = FastAPI()

SLACK_SIGNING_SECRET = "c8f402b38912d04c26478bcb9a8e5049"
SLACK_BOT_TOKEN      = "xoxb-10529798110578-10514391130087-2mNkNg6gIUpJ9lA7rVKjuUiJ"
TARGET_NIF           = "513485279"
DOWNLOAD_DIR         = os.getenv("DOWNLOAD_DIR", "downloads")
OUTPUT_DIR           = os.getenv("OUTPUT_DIR", "outputs")

os.makedirs(DOWNLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Google Sheets
GSHEET_ID          = "1flulATai_w_PxN1e33aokqpqw_wUJqi-2vtQU7bIpOo"
GCREDENTIALS_PATH  = os.getenv("GOOGLE_CREDENTIALS", r"C:\Users\Asus\Downloads\utility-subset-487821-g5-0ff01afdec15.json")
GOAUTH_TOKEN_PATH  = os.getenv("GOOGLE_OAUTH_TOKEN", r"C:\Users\Asus\Downloads\oauth_token.json")
GSCOPES            = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]

# Pasta de fallback (usada se o mês não estiver no dicionário abaixo)
DRIVE_FATURAS_FOLDER_ID = os.getenv("DRIVE_FATURAS_FOLDER", "1RFJ-mkwM7RrHHeujKvFTbDHTJ3FtHOp7")

# IDs das pastas por mês
DRIVE_FOLDER_BY_MONTH = {
    1:  DRIVE_FATURAS_FOLDER_ID,
    2:  DRIVE_FATURAS_FOLDER_ID,
    3:  DRIVE_FATURAS_FOLDER_ID,
    4:  "16E-mugAVyyP4GessoAcEqmlXRupn59b3",
    5:  DRIVE_FATURAS_FOLDER_ID,
    6:  DRIVE_FATURAS_FOLDER_ID,
    7:  DRIVE_FATURAS_FOLDER_ID,
    8:  DRIVE_FATURAS_FOLDER_ID,
    9:  DRIVE_FATURAS_FOLDER_ID,
    10: DRIVE_FATURAS_FOLDER_ID,
    11: DRIVE_FATURAS_FOLDER_ID,
    12: DRIVE_FATURAS_FOLDER_ID,
}

# Cabeçalhos das colunas
SHEET_HEADERS = ["Data", "Evento", "Nome", "Motivo", "IBAN", "Total (€)", "NIF ESN OK?", "Preview", "Fatura Drive", "Pago"]

# =========================
# Deduplicação de eventos
# =========================

_PROCESSED_EVENTS: dict = {}
_EVENT_TTL = 300


def _is_duplicate(event_id: str) -> bool:
    now = time.time()
    for k in list(_PROCESSED_EVENTS):
        if now - _PROCESSED_EVENTS[k] > _EVENT_TTL:
            del _PROCESSED_EVENTS[k]
    if event_id in _PROCESSED_EVENTS:
        return True
    _PROCESSED_EVENTS[event_id] = now
    return False

# =========================
# Slack helpers
# =========================

def verify_slack_signature(signing_secret: str, body: bytes, timestamp: str, signature: str) -> bool:
    try:
        ts = int(timestamp)
    except Exception:
        return False
    if abs(time.time() - ts) > 300:
        return False
    basestring = b"v0:" + timestamp.encode("utf-8") + b":" + body
    my_sig = "v0=" + hmac.new(
        signing_secret.encode("utf-8"), basestring, hashlib.sha256
    ).hexdigest()
    return hmac.compare_digest(my_sig, signature)


def slack_api(method: str, payload: dict) -> dict:
    url = f"https://slack.com/api/{method}"
    headers = {
        "Authorization": f"Bearer {SLACK_BOT_TOKEN}",
        "Content-Type": "application/json; charset=utf-8",
    }
    r = requests.post(url, headers=headers, json=payload, timeout=30)
    r.raise_for_status()
    data = r.json()
    if not data.get("ok"):
        raise RuntimeError(f"Slack API error on {method}: {data}")
    return data


def post_reply(channel: str, thread_ts: str, text: str) -> None:
    slack_api("chat.postMessage", {"channel": channel, "thread_ts": thread_ts, "text": text})


def get_slack_user_real_name(user_id: str) -> Optional[str]:
    """Obtém o nome completo do perfil Slack (real_name) a partir do user_id."""
    try:
        url = "https://slack.com/api/users.info"
        headers = {"Authorization": f"Bearer {SLACK_BOT_TOKEN}"}
        params  = {"user": user_id}
        r = requests.get(url, headers=headers, params=params, timeout=10)
        r.raise_for_status()
        data = r.json()
        if not data.get("ok"):
            return None
        profile = data.get("user", {}).get("profile", {})
        return profile.get("real_name") or profile.get("display_name") or None
    except Exception:
        return None


def get_parent_message_text(channel: str, thread_ts: str) -> str:
    try:
        url = "https://slack.com/api/conversations.replies"
        headers = {"Authorization": f"Bearer {SLACK_BOT_TOKEN}"}
        params  = {"channel": channel, "ts": thread_ts, "limit": 1}
        r = requests.get(url, headers=headers, params=params, timeout=10)
        r.raise_for_status()
        data = r.json()
        if not data.get("ok"):
            return f"[erro API: {data.get('error')}]"
        messages = data.get("messages", [])
        if messages:
            return messages[0].get("text", "")
        return ""
    except Exception as e:
        return f"[erro: {e}]"


def upload_image_to_slack(channel: str, thread_ts: str, image_path: str, title: str) -> None:
    url = "https://slack.com/api/files.upload"
    headers = {"Authorization": f"Bearer {SLACK_BOT_TOKEN}"}
    with open(image_path, "rb") as f:
        resp = requests.post(
            url,
            headers=headers,
            data={
                "channels": channel,
                "thread_ts": thread_ts,
                "title": title,
                "filename": os.path.basename(image_path),
            },
            files={"file": f},
            timeout=60,
        )
    resp.raise_for_status()
    data = resp.json()
    if not data.get("ok"):
        raise RuntimeError(f"Slack files.upload error: {data}")


def download_slack_file(url_private: str, filename: str) -> str:
    local_path = os.path.join(DOWNLOAD_DIR, filename)
    headers = {"Authorization": f"Bearer {SLACK_BOT_TOKEN}"}
    r = requests.get(url_private, headers=headers, stream=True, timeout=60)
    r.raise_for_status()
    with open(local_path, "wb") as f:
        for chunk in r.iter_content(chunk_size=1024 * 1024):
            if chunk:
                f.write(chunk)
    return local_path


def safe_filename(name: str) -> str:
    for ch in r'\/:*?"<>|':
        name = name.replace(ch, "_")
    return name.strip()[:180]

# =========================
# Construção do nome Drive
# =========================

def format_iban_spaced(iban: Optional[str]) -> Optional[str]:
    """
    Converte IBAN sem espaços (PT50001800032531086302097)
    para formato com espaços (PT50 0018 0003 2531 0863 0209 7).
    Devolve None se o IBAN for None.
    """
    if not iban:
        return None
    iban_clean = re.sub(r"\s", "", iban)
    # Agrupa em blocos de 4 (o último pode ter menos)
    return " ".join(iban_clean[i:i+4] for i in range(0, len(iban_clean), 4))


def build_drive_filename(
    nome: Optional[str],
    evento: Optional[str],
    iban: Optional[str],
    ext: str,
    prefix: str = "",
) -> str:
    """
    Constrói o nome do ficheiro para o Drive no formato:
        [PREFIX] Nome - Evento - IBAN.ext
    O IBAN é formatado com espaços para melhor legibilidade.
    Usa fallback incremental se algum campo estiver em falta.
    O prefix é opcional (ex: "PREVIEW_") para distinguir previews das faturas.
    """
    parts = []
    if nome:
        parts.append(nome.strip())
    if evento:
        parts.append(evento.strip())

    iban_fmt = format_iban_spaced(iban)
    if iban_fmt:
        parts.append(iban_fmt)

    if not ext.startswith("."):
        ext = f".{ext}"

    if parts:
        base = " - ".join(parts)
    else:
        base = f"ficheiro_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

    if prefix:
        base = f"{prefix}{base}"

    return safe_filename(base) + ext

# =========================
# Parse message fields
# =========================

def parse_message_fields(text: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Extrai MOTIVO e IBAN da mensagem.
    O NOME já não é lido da mensagem — vem sempre do perfil Slack.
    """
    def pick(label: str) -> Optional[str]:
        m = re.search(rf"(?im)^\s*{label}\s*:\s*(.+?)\s*$", text)
        return m.group(1).strip() if m else None

    motivo = pick("MOTIVO")

    # IBAN: tenta com label primeiro, depois procura o padrão directamente
    iban = pick("IBAN")
    if not iban:
        m = re.search(r"\bPT50\d{21}\b", re.sub(r"\s", "", text))
        if m:
            iban = m.group(0)
        else:
            m = re.search(r"PT\d{2}(?:\s?\d{4}){5,6}", text)
            if m:
                iban = re.sub(r"\s", "", m.group(0))

    return motivo, iban

# =========================
# OCR / PDF text extraction
# =========================

def preprocess_image(img: Image.Image) -> Image.Image:
    img = img.resize((img.width * 2, img.height * 2), Image.LANCZOS)
    img = img.convert("L")
    img = ImageEnhance.Contrast(img).enhance(2.0)
    img = img.filter(ImageFilter.SHARPEN)
    return img


def ocr_image(path: str) -> str:
    img = Image.open(path)
    img = preprocess_image(img)

    attempts = [
        {"lang": "por", "config": "--psm 6"},
        {"lang": "por+eng", "config": "--psm 6"},
        {"config": "--psm 6"},
        {"config": "--psm 4"},
        {"config": "--psm 3"},
    ]

    best = ""
    for kwargs in attempts:
        try:
            t = pytesseract.image_to_string(img, **kwargs)
            if not t or len(t.strip()) < 20:
                continue
            if TARGET_NIF in re.sub(r"\D", "", t):
                return t
            if len(t) > len(best):
                best = t
        except Exception:
            continue

    return best


def extract_text_from_pdf(path: str) -> str:
    doc = fitz.open(path)
    pages = [page.get_text("text") or "" for page in doc]
    return "\n".join(pages).strip()

# =========================
# NIF check
# =========================

def has_target_nif(text: str) -> bool:
    if TARGET_NIF in re.sub(r"\D", "", text):
        return True
    nif_pattern = r"[\s.]?".join(TARGET_NIF)
    return bool(re.search(nif_pattern, text))

# =========================
# Amount parsing
# =========================

def parse_euro_amount(raw: str) -> Optional[float]:
    s = raw.strip()
    s = re.sub(r"[€e\]\|\s]", "", s, flags=re.IGNORECASE)
    if re.search(r"\d\.\d{3},", s):
        s = s.replace(".", "").replace(",", ".")
    elif "," in s and "." not in s:
        s = s.replace(",", ".")
    elif re.fullmatch(r"\d{1,3}\.\d{3}", s):
        s = s.replace(".", "")
    try:
        v = float(s)
        return v if 0 < v < 1_000_000 else None
    except Exception:
        return None

# =========================
# Total extraction
# =========================

_MONEY_WITH_SYMBOL = re.compile(
    r"[€e]\s*(\d[\d\s.,]*\d|\d)\s*[\]\|]?"
    r"|(\d[\d.,]*\d|\d)\s*€",
    re.IGNORECASE,
)
_MONEY_PLAIN = re.compile(
    r"(?<!\d)(\d{1,3}(?:[.,]\d{3})*[.,]\d{2}|\d+[.,]\d{2})(?!\d)",
)
_TOTAL_KEYWORDS = [
    (10, re.compile(r"total\s+geral",         re.IGNORECASE)),
    ( 9, re.compile(r"total\s+a\s+pagar",     re.IGNORECASE)),
    ( 9, re.compile(r"valor\s+a\s+pagar",     re.IGNORECASE)),
    ( 9, re.compile(r"montante\s+a\s+pagar",  re.IGNORECASE)),
    ( 8, re.compile(r"a\s+pagar",             re.IGNORECASE)),
    ( 7, re.compile(r"total\s+documento",     re.IGNORECASE)),
    ( 6, re.compile(r"\btotal\b",             re.IGNORECASE)),
]
_BAD_CONTEXT = re.compile(
    r"\b(subtotal|sub[ -]total|imposto|taxa|desconto|discount|troco|iliquido|incid[eê]ncias)\b"
    r"|\btotal\s+iva\b"
    r"|\bvalor\s+iva\b",
    re.IGNORECASE,
)
_DOC_CONTEXT = re.compile(
    r"\b(ft|fr|fs)\s+\d"
    r"|\batcud\b"
    r"|\bn[ºo]\s*\d",
    re.IGNORECASE,
)


def _extract_amount(snippet: str, allow_plain: bool = False) -> Optional[float]:
    for m in _MONEY_WITH_SYMBOL.finditer(snippet):
        v = parse_euro_amount(m.group(0))
        if v is not None:
            return v
    if allow_plain:
        for m in _MONEY_PLAIN.finditer(snippet):
            v = parse_euro_amount(m.group(0))
            if v is not None:
                return v
    return None


def extract_total_robust(text: str) -> Optional[float]:
    if not text:
        return None

    t = re.sub(r"[ \t]+", " ", text.replace("\xa0", " "))
    lines = t.splitlines()

    expanded = []
    for line in lines:
        if re.search(r"total\s+iva", line, re.IGNORECASE) and \
           re.search(r"total\s+documento", line, re.IGNORECASE):
            for p in re.split(r"(?i)total\s+iva\b", line):
                if p.strip():
                    expanded.append(p.strip())
        else:
            expanded.append(line)
    lines = expanded

    best_score: int = -999
    best_val: Optional[float] = None

    for i, line in enumerate(lines):
        for score, kw_re in _TOTAL_KEYWORDS:
            if not kw_re.search(line):
                continue
            if _BAD_CONTEXT.search(line):
                score -= 7
            if _DOC_CONTEXT.search(line):
                score -= 10
            window = line + " " + " ".join(lines[i + 1: i + 3])
            val = _extract_amount(window, allow_plain=True)
            if val is None:
                continue
            if score > best_score or (score == best_score and val > (best_val or 0)):
                best_score = score
                best_val   = val
            break

    if best_val is not None:
        return best_val

    all_vals = []
    for line in lines:
        if _BAD_CONTEXT.search(line) or _DOC_CONTEXT.search(line):
            continue
        val = _extract_amount(line, allow_plain=True)
        if val and 0 < val < 1_000_000:
            all_vals.append(val)

    return max(all_vals) if all_vals else None

# =========================
# Preview image generator
# =========================

BANNER_BG_OK    = (39, 174, 96)
BANNER_BG_WARN  = (231, 76, 60)
BANNER_TEXT_CLR = (255, 255, 255)
BANNER_HEIGHT   = 90
PREVIEW_MAX_W   = 800


def _get_font(size: int) -> ImageFont.ImageFont:
    for path in [
        r"C:\Windows\Fonts\arialbd.ttf",
        r"C:\Windows\Fonts\arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
    ]:
        if os.path.exists(path):
            try:
                return ImageFont.truetype(path, size)
            except Exception:
                continue
    return ImageFont.load_default()


def generate_preview_image(
    source_path: str,
    total: Optional[float],
    has_nif: bool,
    nome: Optional[str],
    output_path: str,
) -> str:
    ext = os.path.splitext(source_path)[1].lower()

    if ext == ".pdf":
        doc  = fitz.open(source_path)
        pix  = doc[0].get_pixmap(matrix=fitz.Matrix(150/72, 150/72), alpha=False)
        img  = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    else:
        img = Image.open(source_path).convert("RGB")

    if img.width > PREVIEW_MAX_W:
        ratio = PREVIEW_MAX_W / img.width
        img = img.resize((PREVIEW_MAX_W, int(img.height * ratio)), Image.LANCZOS)

    w      = img.width
    banner = Image.new("RGB", (w, BANNER_HEIGHT), BANNER_BG_OK if has_nif else BANNER_BG_WARN)
    draw   = ImageDraw.Draw(banner)

    total_str = f"TOTAL: {total:.2f} €" if total is not None else "TOTAL: NÃO ENCONTRADO"
    nif_str   = "✓ NIF ESN OK" if has_nif else "✗ NIF ESN NÃO ENCONTRADO"
    line2     = nif_str + (f"  |  {nome}" if nome else "")

    draw.text((16, 10), total_str, font=_get_font(32), fill=BANNER_TEXT_CLR)
    draw.text((16, 52), line2,     font=_get_font(18), fill=BANNER_TEXT_CLR)

    combined = Image.new("RGB", (w, BANNER_HEIGHT + img.height), (255, 255, 255))
    combined.paste(banner, (0, 0))
    combined.paste(img,    (0, BANNER_HEIGHT))
    combined.save(output_path, "JPEG", quality=85)
    return output_path

# =========================
# Output (.txt / .docx)
# =========================

def write_summary_files(
    base_name: str,
    nome: Optional[str],
    motivo: Optional[str],
    iban: Optional[str],
    has_nif: bool,
    total: Optional[float],
    source_file: str,
) -> Tuple[str, Optional[str]]:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_base = safe_filename(f"{base_name}_{timestamp}")
    txt_path  = os.path.join(OUTPUT_DIR, f"{safe_base}.txt")
    docx_path = os.path.join(OUTPUT_DIR, f"{safe_base}.docx") if DOCX_AVAILABLE else None

    total_str = f"{total:.2f} €" if total is not None else "NÃO ENCONTRADO"
    nif_str   = f"SIM ({TARGET_NIF})" if has_nif else "NÃO"

    content = (
        f"Tem NIF alvo ({TARGET_NIF})? {nif_str}\n"
        f"Total: {total_str}\n"
        f"Nome: {nome or 'N/D'}\n"
        f"Motivo: {motivo or 'N/D'}\n"
        f"IBAN: {iban or 'N/D'}\n"
        f"Ficheiro origem: {source_file}\n"
    )

    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(content)

    if DOCX_AVAILABLE and docx_path:
        doc = Document()
        doc.add_heading("Resumo da Fatura / Despesa", level=1)
        for line in content.strip().splitlines():
            doc.add_paragraph(line)
        doc.save(docx_path)

    return txt_path, docx_path

# =========================
# Google Sheets
# =========================

def _get_sheets_service():
    creds = Credentials.from_service_account_file(GCREDENTIALS_PATH, scopes=GSCOPES)
    return build("sheets", "v4", credentials=creds, cache_discovery=False)


def _get_drive_service():
    import json as _json
    with open(GOAUTH_TOKEN_PATH, "r") as f:
        token_data = _json.load(f)

    creds = OAuthCredentials(
        token         = token_data.get("token"),
        refresh_token = token_data.get("refresh_token"),
        token_uri     = token_data.get("token_uri", "https://oauth2.googleapis.com/token"),
        client_id     = token_data.get("client_id"),
        client_secret = token_data.get("client_secret"),
        scopes        = token_data.get("scopes"),
    )

    if creds.expired or not creds.valid:
        creds.refresh(GoogleAuthRequest())
        token_data["token"] = creds.token
        with open(GOAUTH_TOKEN_PATH, "w") as f:
            import json as _json2
            _json2.dump(token_data, f, indent=2)

    return build("drive", "v3", credentials=creds, cache_discovery=False)


def _drive_upload(local_path: str, filename: str, mimetype: str, date: datetime = None) -> dict:
    """
    Upload para o Drive na pasta correspondente ao mês da data fornecida.
    O filename deve já vir formatado como 'Nome - Evento - IBAN.ext'.
    """
    from googleapiclient.http import MediaFileUpload

    drive     = _get_drive_service()
    ref_date  = date or datetime.now()
    folder_id = DRIVE_FOLDER_BY_MONTH.get(ref_date.month, DRIVE_FATURAS_FOLDER_ID)

    fmeta = {"name": filename}
    if folder_id:
        fmeta["parents"] = [folder_id]

    media  = MediaFileUpload(local_path, mimetype=mimetype, resumable=False)
    result = drive.files().create(
        body=fmeta,
        media_body=media,
        fields="id,webViewLink",
        supportsAllDrives=True,
    ).execute()

    drive.permissions().create(
        fileId=result["id"],
        body={"type": "anyone", "role": "reader"},
        supportsAllDrives=True,
    ).execute()

    return result


def upload_preview_to_drive(
    image_path: str,
    nome: Optional[str],
    evento: Optional[str],
    iban: Optional[str],
    date: datetime = None,
) -> Optional[str]:
    """Upload do preview com nome 'PREVIEW - Nome - Evento - IBAN.jpg'."""
    filename = build_drive_filename(nome, evento, iban, ext=".jpg", prefix="PREVIEW - ")
    result   = _drive_upload(image_path, filename, "image/jpeg", date=date)
    return f"https://drive.google.com/thumbnail?id={result['id']}&sz=w300"


def upload_fatura_to_drive(
    local_path: str,
    nome: Optional[str],
    evento: Optional[str],
    iban: Optional[str],
    original_ext: str,
    date: datetime = None,
) -> str:
    """Upload da fatura com nome 'Nome - Evento - IBAN.ext'."""
    import mimetypes
    mimetype = mimetypes.guess_type(local_path)[0] or "application/octet-stream"
    filename = build_drive_filename(nome, evento, iban, ext=original_ext)
    result   = _drive_upload(local_path, filename, mimetype, date=date)
    return result.get("webViewLink", f"https://drive.google.com/file/d/{result['id']}/view")


def _get_sheet_name(service) -> str:
    meta = service.spreadsheets().get(spreadsheetId=GSHEET_ID).execute()
    return meta["sheets"][0]["properties"]["title"]


def _ensure_headers(service, sheet_name: str) -> None:
    result = service.spreadsheets().values().get(
        spreadsheetId=GSHEET_ID,
        range=f"{sheet_name}!A1:Z1",
    ).execute()
    if not result.get("values"):
        service.spreadsheets().values().update(
            spreadsheetId=GSHEET_ID,
            range=f"{sheet_name}!A1",
            valueInputOption="RAW",
            body={"values": [SHEET_HEADERS]},
        ).execute()


def append_to_gsheet(
    nome: Optional[str],
    motivo: Optional[str],
    iban: Optional[str],
    total: Optional[float],
    has_nif: bool,
    preview_url: Optional[str],
    fatura_url: Optional[str],
    evento: Optional[str] = None,
) -> None:
    if not GSHEETS_AVAILABLE:
        raise RuntimeError("google-api-python-client não está instalado")
    if not os.path.exists(GCREDENTIALS_PATH):
        raise RuntimeError(f"Credentials não encontradas em: {GCREDENTIALS_PATH}")

    service    = _get_sheets_service()
    sheet_name = _get_sheet_name(service)
    _ensure_headers(service, sheet_name)

    result   = service.spreadsheets().values().get(
        spreadsheetId=GSHEET_ID,
        range=f"{sheet_name}!A:A",
    ).execute()
    next_row = len(result.get("values", [])) + 1

    data_hoje = datetime.now().strftime("%d/%m/%Y %H:%M")
    total_str = f"{total:.2f}" if total is not None else ""
    nif_str   = "SIM" if has_nif else "NÃO"

    if preview_url and "id=" in preview_url:
        img_id        = preview_url.split("id=")[-1].split("&")[0]
        image_formula = f'=IMAGE("https://drive.google.com/thumbnail?id={img_id}&sz=w300")'
    else:
        image_formula = ""

    if fatura_url:
        clean_fatura_url = fatura_url.split("?")[0]
        fatura_link = f'=HYPERLINK("{clean_fatura_url}";"Ver fatura")'
    else:
        fatura_link = ""

    row_formulas = [evento or "", nome or "", motivo or "", iban or "", total_str, nif_str, image_formula, fatura_link, False]

    service.spreadsheets().values().update(
        spreadsheetId=GSHEET_ID,
        range=f"{sheet_name}!A{next_row}",
        valueInputOption="RAW",
        body={"values": [[data_hoje]]},
    ).execute()

    service.spreadsheets().values().update(
        spreadsheetId=GSHEET_ID,
        range=f"{sheet_name}!B{next_row}",
        valueInputOption="USER_ENTERED",
        body={"values": [row_formulas]},
    ).execute()

    sheet_id = service.spreadsheets().get(
        spreadsheetId=GSHEET_ID
    ).execute()["sheets"][0]["properties"]["sheetId"]

    service.spreadsheets().batchUpdate(
        spreadsheetId=GSHEET_ID,
        body={"requests": [
            {
                "setDataValidation": {
                    "range": {
                        "sheetId": sheet_id,
                        "startRowIndex": next_row - 1,
                        "endRowIndex": next_row,
                        "startColumnIndex": 9,
                        "endColumnIndex": 10,
                    },
                    "rule": {"condition": {"type": "BOOLEAN"}, "strict": True},
                }
            },
            {
                "updateDimensionProperties": {
                    "range": {
                        "sheetId": sheet_id,
                        "dimension": "ROWS",
                        "startIndex": next_row - 1,
                        "endIndex": next_row,
                    },
                    "properties": {"pixelSize": 150},
                    "fields": "pixelSize",
                }
            },
        ]},
    ).execute()


# =========================
# Background processor
# =========================

def _process_event(payload: dict) -> None:
    event = payload.get("event", {})

    if event.get("subtype") == "bot_message" or event.get("bot_id"):
        return

    channel   = event.get("channel")
    thread_ts = event.get("ts")
    parent_ts = event.get("thread_ts")
    files     = event.get("files", []) or []

    if not files:
        return

    text = event.get("text", "") or ""

    # --- Nome vem sempre do perfil Slack ---
    user_id = event.get("user")
    nome    = get_slack_user_real_name(user_id) if user_id else None

    # --- IBAN e MOTIVO vêm da mensagem ---
    motivo, iban = parse_message_fields(text)

    # --- Timestamp da mensagem para escolha da pasta Drive ---
    msg_date = datetime.fromtimestamp(float(event.get("ts", 0)))

    # --- Evento: lê a mensagem pai da thread ---
    evento = None
    if parent_ts and parent_ts != thread_ts:
        parent_text = get_parent_message_text(channel, parent_ts)
        m = re.search(r"(?im)^\s*EVENTO\s*:\s*(.+?)\s*$", parent_text)
        if m:
            evento = m.group(1).strip()
        else:
            first_line = parent_text.strip().splitlines()[0].strip() if parent_text.strip() else None
            if first_line and len(first_line) < 100:
                evento = first_line
        thread_ts = parent_ts

    # Avisa se IBAN estiver em falta
    if not iban:
        post_reply(
            channel, thread_ts,
            f"⚠️ Não encontrei um IBAN na tua mensagem.\n"
            f"Envia o IBAN no formato: `PT50XXXXXXXXXXXXXXXXXXX`\n"
            f"A processar o ficheiro com os dados disponíveis..."
        )

    for f in files:
        mimetype      = f.get("mimetype", "")
        original_name = f.get("name") or f.get("title") or f.get("id") or "ficheiro"
        original_name = safe_filename(original_name)
        original_ext  = os.path.splitext(original_name)[1]  # ex: .pdf, .jpg
        url_private   = f.get("url_private_download") or f.get("url_private")

        if not url_private:
            continue

        # 1) Download
        try:
            local_path = download_slack_file(url_private, original_name)
        except Exception as e:
            post_reply(channel, thread_ts, f"⚠️ Erro ao descarregar {original_name}: {e}")
            continue

        # 1b) Upload da fatura para o Drive com nome "Nome - Evento - IBAN.ext"
        fatura_drive_url = None
        try:
            fatura_drive_url = upload_fatura_to_drive(
                local_path   = local_path,
                nome         = nome,
                evento       = evento,
                iban         = iban,
                original_ext = original_ext,
                date         = msg_date,
            )
        except Exception as e:
            post_reply(channel, thread_ts, f"⚠️ Fatura Drive falhou: {e}")

        # 2) Extrair texto
        extracted_text = ""
        ext = os.path.splitext(local_path)[1].lower()
        try:
            if ext == ".pdf":
                extracted_text = extract_text_from_pdf(local_path)
            elif mimetype.startswith("image/") or ext in (".jpg", ".jpeg", ".png", ".webp", ".tif", ".tiff"):
                extracted_text = ocr_image(local_path)
        except Exception as e:
            post_reply(channel, thread_ts, f"⚠️ Erro a extrair texto de {original_name}: {e}")

        # DEBUG: guarda OCR em ficheiro
        try:
            dbg = os.path.join(OUTPUT_DIR, f"DEBUG_OCR_{safe_filename(original_name)}.txt")
            with open(dbg, "w", encoding="utf-8") as dbf:
                dbf.write(f"=== OCR TEXT ===\n{extracted_text}\n\n")
                dbf.write(f"=== ALL DIGITS ===\n{re.sub(r'[^0-9]', '', extracted_text)}\n\n")
                dbf.write(f"=== NIF FOUND ===\n{has_target_nif(extracted_text)}\n")
        except Exception:
            pass

        # 3) Regras
        found_nif = has_target_nif(extracted_text) if extracted_text else False
        max_total = extract_total_robust(extracted_text) if extracted_text else None

        # 4) Preview com banner — guardado localmente com nome formatado
        preview_path = None
        preview_url  = None
        try:
            preview_filename = build_drive_filename(nome, evento, iban, ext=".jpg", prefix="PREVIEW - ")
            preview_path     = os.path.join(OUTPUT_DIR, preview_filename)
            generate_preview_image(local_path, max_total, found_nif, nome, preview_path)
        except Exception as e:
            preview_path = None
            post_reply(channel, thread_ts, f"⚠️ Não foi possível gerar preview: {e}")

        # 4b) Upload do preview ao Drive com nome "PREVIEW - Nome - Evento - IBAN.jpg"
        if preview_path and os.path.exists(preview_path):
            try:
                preview_url = upload_preview_to_drive(
                    image_path = preview_path,
                    nome       = nome,
                    evento     = evento,
                    iban       = iban,
                    date       = msg_date,
                )
            except Exception as e:
                preview_url = None
                post_reply(channel, thread_ts, f"⚠️ Preview Drive falhou: {e}")

        # 5) Resumo TXT/DOCX
        txt_path, docx_path = write_summary_files(
            base_name   = os.path.splitext(original_name)[0],
            nome        = nome,
            motivo      = motivo,
            iban        = iban,
            has_nif     = found_nif,
            total       = max_total,
            source_file = local_path,
        )

        # 6) Escreve no Google Sheet
        sheet_ok  = False
        sheet_err = ""
        try:
            append_to_gsheet(
                nome        = nome,
                motivo      = motivo,
                iban        = iban,
                total       = max_total,
                has_nif     = found_nif,
                preview_url = preview_url,
                fatura_url  = fatura_drive_url,
                evento      = evento,
            )
            sheet_ok = True
        except Exception as e:
            sheet_err = str(e)

        # 7) Confirmação no Slack
        total_str    = f"{max_total:.2f} €" if max_total is not None else "NÃO ENCONTRADO"
        nif_str      = "SIM ✅" if found_nif else "NÃO ❌"
        mes_str      = msg_date.strftime("%B %Y")
        drive_name   = build_drive_filename(nome, evento, iban, ext=original_ext)

        if sheet_ok:
            msg = (
                f"✅ *{drive_name}* registado no sheet!\n"
                f"• Total: *{total_str}*  |  NIF ESN: *{nif_str}*\n"
                f"• Nome: {nome or 'N/D'}  |  Motivo: {motivo or 'N/D'}\n"
                f"• Pasta Drive: *{mes_str}*"
            )
        else:
            msg = (
                f"📄 *{drive_name}* processado (⚠️ sheet falhou: {sheet_err})\n"
                f"• NIF alvo ({TARGET_NIF})? *{nif_str}*\n"
                f"• Total: *{total_str}*\n"
                f"• Nome: {nome or 'N/D'}\n"
                f"• Motivo: {motivo or 'N/D'}\n"
                f"• IBAN: {iban or 'N/D'}\n"
                f"• Pasta Drive: *{mes_str}*"
            )

        post_reply(channel, thread_ts, msg)

        # 8) Preview como imagem na thread do Slack
        if preview_path and os.path.exists(preview_path):
            try:
                upload_image_to_slack(channel, thread_ts, preview_path, f"Preview — {drive_name}")
            except Exception as e:
                post_reply(channel, thread_ts, f"⚠️ Erro ao enviar preview: {e}")


# =========================
# Endpoint principal
# =========================

@app.post("/slack/events")
async def slack_events(request: Request):
    body      = await request.body()
    timestamp = request.headers.get("X-Slack-Request-Timestamp", "")
    signature = request.headers.get("X-Slack-Signature", "")

    if not verify_slack_signature(SLACK_SIGNING_SECRET, body, timestamp, signature):
        raise HTTPException(status_code=401, detail="Invalid Slack signature")

    payload = json.loads(body.decode("utf-8"))

    if payload.get("type") == "url_verification":
        return JSONResponse({"challenge": payload.get("challenge")})

    if payload.get("type") != "event_callback":
        return JSONResponse({"ok": True})

    event_id = payload.get("event_id", "")
    if event_id and _is_duplicate(event_id):
        return JSONResponse({"ok": True})

    Thread(target=_process_event, args=(payload,), daemon=True).start()

    return JSONResponse({"ok": True})